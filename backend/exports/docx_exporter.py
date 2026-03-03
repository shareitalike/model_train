from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


class DocxExporter:
    def export(self, result_data: dict) -> bytes:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        doc.core_properties.title = "कैथी OCR आउटपुट"
        doc.core_properties.author = "Kaithi Digitization System"

        title = doc.add_heading("कैथी लिपि → हिन्दी रूपान्तरण", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)

        meta = result_data.get("metadata", {})
        conf = result_data.get("overall_confidence", 0)
        tbl = doc.add_table(rows=2, cols=3)
        tbl.style = "Table Grid"
        for i, h in enumerate(["कुल पृष्ठ", "विश्वास स्तर", "क्षेत्र"]):
            tbl.rows[0].cells[i].text = h
        for i, v in enumerate([
            str(meta.get("total_pages", "N/A")),
            f"{conf*100:.1f}%",
            meta.get("region", "मानक"),
        ]):
            tbl.rows[1].cells[i].text = v
        doc.add_paragraph()

        for page in result_data.get("pages", []):
            h = doc.add_heading(f"पृष्ठ {page.get('page_number', '?')}", 2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
            hindi = page.get("corrected_text") or page.get("hindi_text", "")
            if hindi:
                para = doc.add_paragraph(hindi)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = "Noto Sans Devanagari"
                    run.font.size = Pt(14)
            note = doc.add_paragraph(
                f"विश्वास: {page.get('confidence', 0)*100:.1f}% | "
                f"पंक्तियाँ: {page.get('line_count', 0)}"
            )
            for run in note.runs:
                run.font.size = Pt(9)
            doc.add_page_break()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
